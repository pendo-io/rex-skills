"""
generate_scoping_doc.py

Generates a Pendo-branded Agent/Module Scoping Checklist .docx by appending
filled checklist content to the canonical Pendo template. Applies Pendo Pink
(#FF4081) to the document title and all bordered table header rows.

Usage:
    from generate_scoping_doc import generate_scoping_doc

    output_path = generate_scoping_doc(
        template_path="/home/claude/template.docx",
        output_path="/mnt/user-data/outputs/MyProject_Scoping_Checklist.docx",
        project_name="My Project",
        project_description="One or two sentence description.",
        modules=[...],           # see MODULE_SCHEMA below
        cc_cd_framework={...},   # optional, see CC_CD_SCHEMA below
        generated_date="May 14, 2026",
    )

MODULE_SCHEMA:
    {
        "name": str,
        "jtbd": str,
        "trigger": str,
        "agent_type": {
            "selected": "Action Driver | Intelligence Provider | Autonomous",
            "rationale": str,
        },
        "data_sources": [
            {"system": str, "what": str, "access_pattern": str},
        ],
        "output": str,
        "actions": str,
        "suppression": [str, ...],
        "success_metrics": [
            {"type": "Outcome | Adoption | Experience", "metric": str, "target": str},
        ],
        "build_vs_buy": str,
    }

CC_CD_SCHEMA (optional):
    {
        "phases": [
            ["V1 (Launch)", "High", "Low", "Status description"],
            ["V2", "Medium", "Medium", "Status description"],
            ["V3", "Low", "High", "Status description"],
        ],
        "evaluation": str,
        "kpis": [
            ["Module name", "Business KPI", "UX KPI"],
        ],
        "ownership": str,
    }
"""

from __future__ import annotations

from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PENDO_PINK = RGBColor(0xFF, 0x40, 0x81)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Brand styling helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_fill: str = "FF4081") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _style_header_row(row) -> None:
    for cell in row.cells:
        _shade_cell(cell, "FF4081")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
            if not paragraph.runs:
                run = paragraph.add_run()
                run.font.color.rgb = WHITE
                run.font.bold = True


def _pink_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = PENDO_PINK


# ---------------------------------------------------------------------------
# Table builders
# ---------------------------------------------------------------------------

def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = val
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True


def _add_header_table(doc: Document, headers: list[str], data: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    _style_header_row(header_row)
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = val


# ---------------------------------------------------------------------------
# Module section writers
# ---------------------------------------------------------------------------

def _write_module(doc: Document, module: dict[str, Any], idx: int) -> None:
    module_label = module.get("name") or f"Module {idx}"
    _pink_heading(doc, f"Module {idx}: {module_label}", level=2)

    # 1. JTBD
    _pink_heading(doc, "1. Job to Be Done (JTBD)", level=3)
    doc.add_paragraph(module.get("jtbd", "TBD"))

    # 2. Trigger
    _pink_heading(doc, "2. Trigger", level=3)
    doc.add_paragraph(module.get("trigger", "TBD"))

    # 3. Experience Type
    _pink_heading(doc, "3. Experience Type", level=3)
    agent_type = module.get("agent_type", {})
    _add_header_table(
        doc,
        ["Type", "Selected", "Rationale"],
        [
            [
                agent_type.get("selected", "TBD"),
                "Yes",
                agent_type.get("rationale", ""),
            ]
        ],
    )
    doc.add_paragraph()

    # 4. Data Sources
    _pink_heading(doc, "4. Data Sources Required", level=3)
    sources = module.get("data_sources", [])
    if sources:
        _add_header_table(
            doc,
            ["System", "What", "Access Pattern"],
            [[s.get("system", ""), s.get("what", ""), s.get("access_pattern", "")] for s in sources],
        )
    else:
        doc.add_paragraph("TBD")
    doc.add_paragraph()

    # 5. Output
    _pink_heading(doc, "5. Output", level=3)
    doc.add_paragraph(module.get("output", "TBD"))

    # 6. Actions
    _pink_heading(doc, "6. Action(s)", level=3)
    doc.add_paragraph(module.get("actions", "TBD"))

    # 7. Suppression
    _pink_heading(doc, "7. Suppression & Override Logic", level=3)
    suppression = module.get("suppression", [])
    if suppression:
        for rule in suppression:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(rule)
    else:
        doc.add_paragraph("TBD")

    # 8. Success Metrics
    _pink_heading(doc, "8. Success Metrics", level=3)
    metrics = module.get("success_metrics", [])
    if metrics:
        _add_header_table(
            doc,
            ["Type", "Metric", "Target"],
            [[m.get("type", ""), m.get("metric", ""), m.get("target", "")] for m in metrics],
        )
    else:
        doc.add_paragraph("TBD")
    doc.add_paragraph()

    # 9. Build vs. Buy
    _pink_heading(doc, "9. Build vs. Buy", level=3)
    doc.add_paragraph(module.get("build_vs_buy", "TBD"))


# ---------------------------------------------------------------------------
# CC/CD framework writer
# ---------------------------------------------------------------------------

def _write_cc_cd(doc: Document, cc_cd: dict[str, Any]) -> None:
    doc.add_page_break()
    _pink_heading(doc, "Post-Launch: CC/CD Framework", level=2)

    # Calibration phases
    _pink_heading(doc, "1. Calibration Loop", level=3)
    phases = cc_cd.get("phases", [])
    if phases:
        _add_header_table(
            doc,
            ["Phase", "Control Level", "Agency Level", "Status"],
            phases,
        )
    doc.add_paragraph()

    # Evaluation
    _pink_heading(doc, "2. Evaluation Plan", level=3)
    doc.add_paragraph(cc_cd.get("evaluation", "TBD"))

    # KPI table
    _pink_heading(doc, "3. Business KPIs vs UX KPIs", level=3)
    kpis = cc_cd.get("kpis", [])
    if kpis:
        _add_header_table(doc, ["Module", "Business KPI", "UX KPI"], kpis)
    doc.add_paragraph()

    # Ownership
    _pink_heading(doc, "4. Collaboration & Ownership", level=3)
    doc.add_paragraph(cc_cd.get("ownership", "TBD"))


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_scoping_doc(
    template_path: str | Path,
    output_path: str | Path,
    project_name: str,
    project_description: str,
    modules: list[dict[str, Any]],
    cc_cd_framework: dict[str, Any] | None = None,
    generated_date: str | None = None,
) -> str:
    """
    Generate a Pendo-branded scoping checklist .docx.

    Returns the output_path as a string.
    """
    generated_date = generated_date or date.today().strftime("%B %d, %Y")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    # Page break to separate from template content
    doc.add_page_break()

    # Document title (Pendo Pink)
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{project_name} — Agent Scoping Checklist")
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = PENDO_PINK
    title_run.font.bold = True

    doc.add_paragraph(f"Generated: {generated_date}")
    doc.add_paragraph(project_description)
    doc.add_paragraph()

    # Write each module
    for i, module in enumerate(modules, start=1):
        _write_module(doc, module, i)
        if i < len(modules):
            doc.add_paragraph()

    # Optional CC/CD framework
    if cc_cd_framework:
        _write_cc_cd(doc, cc_cd_framework)

    doc.save(str(output_path))
    return str(output_path)
